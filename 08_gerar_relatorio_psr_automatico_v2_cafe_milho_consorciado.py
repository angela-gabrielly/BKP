# -*- coding: utf-8 -*-
"""
08_gerar_relatorio_psr_automatico.py

Gera automaticamente:
1. imagens de mapas, gráficos e tabelas a partir dos GeoPackages do projeto PSR;
2. um relatório Word (.docx) já com as imagens inseridas.

Pré-requisitos esperados:
- 07b_agregar_psr_nas_grades.py
- 07c_agregar_psr_grade_cultura_seguradora.py
- 07d_agregar_psr_grade_cultura_seguradora_sem_dominantes.py (opcional, mas recomendado)

Pacotes:
    pip install geopandas pyogrio matplotlib python-docx pandas numpy

Observações metodológicas:
- Os mapas de pontos/hexágonos representam coordenadas informadas na base PSR;
  não representam o limite real das propriedades.
- Os valores de indenização são valores registrados na base, não uma
  sinistralidade técnica definitiva.
"""

from __future__ import annotations

from pathlib import Path
from datetime import datetime
import math
import textwrap

import geopandas as gpd
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# =============================================================================
# CONFIGURAÇÃO
# =============================================================================

PASTA_PROJETO = Path(r"D:\AZ\PSR")

PASTA_VETORES = PASTA_PROJETO / "data" / "vetores"
PASTA_SAIDA = PASTA_PROJETO / "relatorio_automatico"
PASTA_IMAGENS = PASTA_SAIDA / "imagens"
PASTA_TABELAS = PASTA_SAIDA / "tabelas"

for pasta in [PASTA_SAIDA, PASTA_IMAGENS, PASTA_TABELAS]:
    pasta.mkdir(parents=True, exist_ok=True)

# -----------------------------------------------------------------------------
# GeoPackages gerados nos scripts anteriores.
# -----------------------------------------------------------------------------
GPKG_TOTAL = PASTA_VETORES / "psr_grades_2020_2024.gpkg"
GPKG_CULTURA = PASTA_VETORES / "psr_grades_culturas_2020_2024.gpkg"
GPKG_SEM_DOMINANTES = (
    PASTA_VETORES
    / "psr_grades_culturas_sem_dominantes_2020_2024.gpkg"
)

# Opcional: informe uma malha de municípios/UF para sobrepor nos mapas.
# Exemplo: Path(r"D:\AZ\PSR\data\referencias\BR_Municipios_2024.shp")
# Mantenha None para gerar apenas os hexágonos.
CAMINHO_MALHA_CONTORNO = None
CAMADA_MALHA_CONTORNO = None

# Escolha a escala espacial principal do relatório.
PREFIXO_GRADE = "h10"

# Culturas que serão destacadas no relatório.
CULTURAS_DESTAQUE = [
    "SOJA",
    "MILHO 1ª SAFRA",
    "MILHO 2ª SAFRA",
    "MAÇÃ",
    "UVA",
    "CAFÉ",
]

# A cultura usada para os mapas comparativos de concorrência.
CULTURA_MAPA_CONCORRENCIA = "SOJA"

# Quantidade de linhas para rankings e tabelas no Word.
TOP_N = 10

# Tamanho das imagens em polegadas.
LARGURA_IMAGEM_WORD = Inches(6.25)

# Nome final do relatório.
ARQUIVO_DOCX = PASTA_SAIDA / "Relatorio_Automatico_Inteligencia_PSR_2020_2024.docx"

# Para reprocessar imagens e Word do zero.
LIMPAR_IMAGENS_ANTERIORES = True


# =============================================================================
# FUNÇÕES DE LEITURA E PREPARAÇÃO
# =============================================================================


def verificar_arquivo(caminho: Path, obrigatorio: bool = True) -> bool:
    """Valida se um arquivo existe."""
    existe = caminho.exists()
    if obrigatorio and not existe:
        raise FileNotFoundError(
            f"Arquivo não encontrado:\n{caminho}\n\n"
            "Execute os scripts de geração das grades antes do relatório."
        )
    return existe


def ler_camada(caminho: Path, camada: str) -> gpd.GeoDataFrame:
    """Lê uma camada do GeoPackage e retorna um GeoDataFrame."""
    try:
        gdf = gpd.read_file(caminho, layer=camada, engine="pyogrio")
    except Exception as exc:
        raise RuntimeError(
            f"Não foi possível ler a camada '{camada}' em:\n{caminho}\n\n"
            f"Detalhe: {exc}"
        ) from exc

    if gdf.empty:
        raise ValueError(
            f"A camada '{camada}' está vazia: {caminho}"
        )

    return gdf


def carregar_contorno() -> gpd.GeoDataFrame | None:
    """Carrega contorno opcional para contextualizar os mapas."""
    if CAMINHO_MALHA_CONTORNO is None:
        return None

    if not CAMINHO_MALHA_CONTORNO.exists():
        print(
            "Aviso: contorno não encontrado. Os mapas serão gerados sem sobreposição."
        )
        return None

    try:
        gdf = gpd.read_file(
            CAMINHO_MALHA_CONTORNO,
            layer=CAMADA_MALHA_CONTORNO,
            engine="pyogrio",
        )
    except Exception as exc:
        print(f"Aviso: falha ao ler contorno: {exc}")
        return None

    return gdf


def limpar_nome_arquivo(texto: str) -> str:
    """Cria nomes de arquivo simples e seguros."""
    substituicoes = {
        "Á": "A", "À": "A", "Ã": "A", "Â": "A",
        "É": "E", "Ê": "E", "Í": "I",
        "Ó": "O", "Ô": "O", "Õ": "O",
        "Ú": "U", "Ç": "C", "ª": "A", "º": "O",
        " ": "_", "/": "_", "-": "_",
    }

    resultado = str(texto).upper()
    for antigo, novo in substituicoes.items():
        resultado = resultado.replace(antigo, novo)

    return "".join(
        caractere for caractere in resultado
        if caractere.isalnum() or caractere == "_"
    )



def normalizar_chave_cultura(serie: pd.Series) -> pd.Series:
    """Normaliza cultura para comparação sem acentos e variações de espaçamento."""
    import unicodedata

    def _normalizar(valor) -> str:
        if pd.isna(valor):
            return ""
        texto = unicodedata.normalize("NFKD", str(valor))
        texto = "".join(c for c in texto if not unicodedata.combining(c))
        return " ".join(texto.upper().strip().split())

    return serie.astype("string").map(_normalizar)


def padronizar_cultura(serie: pd.Series) -> pd.Series:
    """
    Padroniza culturas prioritárias para o relatório.

    Regra: MILHO CONSORCIADO e variações são somados a MILHO 2ª SAFRA.
    """
    resultado = serie.astype("string").str.strip().str.upper()
    chave = normalizar_chave_cultura(resultado)

    milho_consorciado = (
        chave.str.contains("MILHO", na=False)
        & chave.str.contains("CONSORCI", na=False)
    )
    milho_1 = (
        chave.str.contains("MILHO", na=False)
        & (
            chave.str.contains(r"\b1\b", regex=True, na=False)
            | chave.str.contains("PRIMEIRA", na=False)
        )
        & chave.str.contains("SAFRA", na=False)
    )
    milho_2 = (
        chave.str.contains("MILHO", na=False)
        & (
            chave.str.contains(r"\b2\b", regex=True, na=False)
            | chave.str.contains("SEGUNDA", na=False)
        )
        & chave.str.contains("SAFRA", na=False)
    )

    resultado.loc[chave.str.startswith("SOJA", na=False)] = "SOJA"
    resultado.loc[milho_1] = "MILHO 1ª SAFRA"
    resultado.loc[milho_2 | milho_consorciado] = "MILHO 2ª SAFRA"
    resultado.loc[chave.str.startswith("MACA", na=False)] = "MAÇÃ"
    resultado.loc[chave.str.startswith("UVA", na=False)] = "UVA"
    resultado.loc[chave.str.startswith("CAFE", na=False)] = "CAFÉ"

    return resultado


def reagregar_camada_cultura_para_relatorio(
    gdf_cult: gpd.GeoDataFrame,
    gdf_cultseg: gpd.GeoDataFrame,
) -> gpd.GeoDataFrame:
    """
    Reagrega a camada hexágono+cultura após a padronização de culturas.

    Isso é necessário para que registros originalmente classificados como
    'MILHO CONSORCIADO' sejam efetivamente somados a 'MILHO 2ª SAFRA'
    nos mapas, tabelas, líder de mercado e HHI.
    """
    gdf_cult = gdf_cult.copy()
    gdf_cultseg = gdf_cultseg.copy()

    gdf_cult["cultura"] = padronizar_cultura(gdf_cult["cultura"])
    gdf_cultseg["cultura"] = padronizar_cultura(gdf_cultseg["cultura"])

    # Base geral por hexágono+cultura. Mantém geometria da célula.
    colunas_soma = [
        "n_registros", "n_localizacoes", "area_total_ha",
        "limite_garantia_rs", "premio_liq_rs", "subvencao_rs",
        "indenizacao_rs", "n_registros_indenizados",
    ]
    for coluna in colunas_soma:
        if coluna not in gdf_cult.columns:
            gdf_cult[coluna] = 0
        gdf_cult[coluna] = pd.to_numeric(gdf_cult[coluna], errors="coerce").fillna(0)

    base = (
        gdf_cult.groupby(["grid_id", "cultura"], as_index=False)
        .agg(
            geometry=("geometry", "first"),
            n_registros=("n_registros", "sum"),
            n_localizacoes=("n_localizacoes", "sum"),
            area_total_ha=("area_total_ha", "sum"),
            limite_garantia_rs=("limite_garantia_rs", "sum"),
            premio_liq_rs=("premio_liq_rs", "sum"),
            subvencao_rs=("subvencao_rs", "sum"),
            indenizacao_rs=("indenizacao_rs", "sum"),
            n_registros_indenizados=("n_registros_indenizados", "sum"),
        )
    )

    # Recalcula concorrência a partir da camada detalhada hex+cultura+seguradora.
    for coluna in [
        "premio_liq_rs", "limite_garantia_rs", "indenizacao_rs",
        "area_total_ha", "subvencao_rs", "n_registros", "n_localizacoes",
        "n_registros_indenizados",
    ]:
        if coluna not in gdf_cultseg.columns:
            gdf_cultseg[coluna] = 0
        gdf_cultseg[coluna] = pd.to_numeric(
            gdf_cultseg[coluna], errors="coerce"
        ).fillna(0)

    cultseg = (
        gdf_cultseg.groupby(["grid_id", "cultura", "seguradora"], as_index=False)
        .agg(
            premio_liq_rs=("premio_liq_rs", "sum"),
            limite_garantia_rs=("limite_garantia_rs", "sum"),
            indenizacao_rs=("indenizacao_rs", "sum"),
            area_total_ha=("area_total_ha", "sum"),
            subvencao_rs=("subvencao_rs", "sum"),
            n_registros=("n_registros", "sum"),
            n_localizacoes=("n_localizacoes", "sum"),
            n_registros_indenizados=("n_registros_indenizados", "sum"),
        )
    )

    chaves = ["grid_id", "cultura"]
    cultseg["premio_total_cultura_hex"] = (
        cultseg.groupby(chaves)["premio_liq_rs"].transform("sum")
    )
    cultseg["market_share_seg_pct"] = np.where(
        cultseg["premio_total_cultura_hex"] > 0,
        cultseg["premio_liq_rs"] / cultseg["premio_total_cultura_hex"] * 100,
        np.nan,
    )
    cultseg["ranking_premio_seg"] = (
        cultseg.groupby(chaves)["premio_liq_rs"]
        .rank(method="dense", ascending=False)
        .astype("Int64")
    )
    cultseg["eh_seguradora_lider"] = cultseg["ranking_premio_seg"].eq(1).astype("int8")

    ranking = cultseg.sort_values(
        chaves + ["premio_liq_rs", "seguradora"],
        ascending=[True, True, False, True],
    )
    lider = (
        ranking.drop_duplicates(subset=chaves, keep="first")
        [chaves + ["seguradora", "market_share_seg_pct"]]
        .rename(columns={
            "seguradora": "seguradora_lider",
            "market_share_seg_pct": "share_seguradora_lider_pct",
        })
    )

    hhi = cultseg.assign(
        hhi_parcela=(cultseg["market_share_seg_pct"].fillna(0) / 100) ** 2
    ).groupby(chaves, as_index=False)["hhi_parcela"].sum()
    hhi["hhi_premio"] = hhi["hhi_parcela"] * 10000
    hhi["classificacao_concentracao"] = np.select(
        [hhi["hhi_premio"] < 1500, hhi["hhi_premio"] < 2500],
        ["BAIXA_CONCENTRACAO", "CONCENTRACAO_MODERADA"],
        default="ALTA_CONCENTRACAO",
    )
    hhi = hhi[chaves + ["hhi_premio", "classificacao_concentracao"]]

    n_seguradoras = (
        cultseg.groupby(chaves, as_index=False)["seguradora"]
        .nunique()
        .rename(columns={"seguradora": "n_seguradoras"})
    )

    base = (
        base.merge(lider, on=chaves, how="left")
        .merge(hhi, on=chaves, how="left")
        .merge(n_seguradoras, on=chaves, how="left")
    )

    base["freq_indenizacao_pct"] = np.where(
        base["n_registros"] > 0,
        base["n_registros_indenizados"] / base["n_registros"] * 100,
        np.nan,
    )
    base["relacao_indenizacao_premio_pct"] = np.where(
        base["premio_liq_rs"] > 0,
        base["indenizacao_rs"] / base["premio_liq_rs"] * 100,
        np.nan,
    )
    base["relacao_indenizacao_limite_pct"] = np.where(
        base["limite_garantia_rs"] > 0,
        base["indenizacao_rs"] / base["limite_garantia_rs"] * 100,
        np.nan,
    )
    base["premio_por_ha"] = np.where(
        base["area_total_ha"] > 0,
        base["premio_liq_rs"] / base["area_total_ha"],
        np.nan,
    )
    base["limite_por_ha"] = np.where(
        base["area_total_ha"] > 0,
        base["limite_garantia_rs"] / base["area_total_ha"],
        np.nan,
    )

    crs = gdf_cult.crs
    return (
        gpd.GeoDataFrame(base, geometry="geometry", crs=crs),
        gpd.GeoDataFrame(
            cultseg.merge(
                gdf_cult[["grid_id", "geometry"]].drop_duplicates("grid_id"),
                on="grid_id", how="left"
            ),
            geometry="geometry", crs=crs
        ),
    )


def formatar_brl(valor: float | int | None) -> str:
    """Formata valores monetários em reais."""
    if valor is None or pd.isna(valor):
        return "-"

    return (
        f"R$ {float(valor):,.0f}"
        .replace(",", "X")
        .replace(".", ",")
        .replace("X", ".")
    )


def formatar_numero(valor: float | int | None, casas: int = 0) -> str:
    """Formata número no padrão brasileiro."""
    if valor is None or pd.isna(valor):
        return "-"

    return (
        f"{float(valor):,.{casas}f}"
        .replace(",", "X")
        .replace(".", ",")
        .replace("X", ".")
    )


def formatar_pct(valor: float | int | None, casas: int = 1) -> str:
    """Formata percentual."""
    if valor is None or pd.isna(valor):
        return "-"
    return f"{formatar_numero(valor, casas)}%"


def garantir_coluna(gdf: pd.DataFrame, coluna: str) -> None:
    """Interrompe com mensagem amigável quando campo necessário não existe."""
    if coluna not in gdf.columns:
        raise KeyError(
            f"Campo obrigatório '{coluna}' não encontrado.\n"
            f"Campos disponíveis: {list(gdf.columns)}"
        )


# =============================================================================
# FUNÇÕES DE MAPAS E GRÁFICOS
# =============================================================================


def configurar_eixos_mapa(ax: plt.Axes, titulo: str) -> None:
    """Padroniza aparência dos mapas."""
    ax.set_title(titulo, fontsize=13, fontweight="bold", pad=12)
    ax.set_axis_off()


def plotar_contorno(ax: plt.Axes, contorno: gpd.GeoDataFrame | None, crs_destino) -> None:
    """Sobrepõe contorno opcional nos mapas."""
    if contorno is None:
        return

    try:
        contorno_plot = contorno.to_crs(crs_destino)
        contorno_plot.boundary.plot(ax=ax, linewidth=0.25, color="gray", alpha=0.55)
    except Exception as exc:
        print(f"Aviso: não foi possível sobrepor contorno: {exc}")


def mapa_graduado(
    gdf: gpd.GeoDataFrame,
    coluna: str,
    titulo: str,
    arquivo_saida: Path,
    contorno: gpd.GeoDataFrame | None = None,
    legenda: str | None = None,
    classes: int = 6,
) -> Path:
    """Gera mapa graduado para um atributo numérico."""
    garantir_coluna(gdf, coluna)

    dados = gdf.copy()
    dados[coluna] = pd.to_numeric(dados[coluna], errors="coerce")
    dados = dados[dados[coluna].notna()].copy()

    if dados.empty:
        raise ValueError(f"Sem valores válidos para mapear '{coluna}'.")

    fig, ax = plt.subplots(figsize=(10, 7.2), dpi=180)
    plotar_contorno(ax, contorno, dados.crs)

    # quantiles lida melhor com distribuição altamente assimétrica de prêmio/limite.
    dados.plot(
        ax=ax,
        column=coluna,
        cmap="YlOrRd",
        scheme="Quantiles",
        k=classes,
        legend=True,
        legend_kwds={
            "title": legenda or coluna,
            "loc": "lower left",
            "fmt": "{:.0f}",
        },
        linewidth=0.10,
        edgecolor="white",
        alpha=0.92,
    )

    configurar_eixos_mapa(ax, titulo)
    fig.tight_layout()
    fig.savefig(arquivo_saida, bbox_inches="tight")
    plt.close(fig)
    return arquivo_saida


def mapa_categorico(
    gdf: gpd.GeoDataFrame,
    coluna: str,
    titulo: str,
    arquivo_saida: Path,
    contorno: gpd.GeoDataFrame | None = None,
    max_categorias: int = 12,
) -> Path:
    """Gera mapa categorizado, agrupando excesso de categorias em OUTROS."""
    garantir_coluna(gdf, coluna)

    dados = gdf.copy()
    dados[coluna] = dados[coluna].fillna("SEM INFORMACAO").astype(str)

    frequencia = dados[coluna].value_counts()
    categorias_mantidas = list(frequencia.head(max_categorias - 1).index)

    dados["categoria_mapa"] = np.where(
        dados[coluna].isin(categorias_mantidas),
        dados[coluna],
        "OUTROS",
    )

    fig, ax = plt.subplots(figsize=(10, 7.2), dpi=180)
    plotar_contorno(ax, contorno, dados.crs)

    dados.plot(
        ax=ax,
        column="categoria_mapa",
        categorical=True,
        legend=True,
        legend_kwds={"title": coluna, "loc": "lower left"},
        linewidth=0.10,
        edgecolor="white",
        alpha=0.92,
    )

    configurar_eixos_mapa(ax, titulo)
    fig.tight_layout()
    fig.savefig(arquivo_saida, bbox_inches="tight")
    plt.close(fig)
    return arquivo_saida


def grafico_barras_horizontais(
    df: pd.DataFrame,
    coluna_categoria: str,
    coluna_valor: str,
    titulo: str,
    arquivo_saida: Path,
    top_n: int = TOP_N,
    rotulo_monetario: bool = False,
) -> Path:
    """Gera ranking horizontal."""
    dados = df[[coluna_categoria, coluna_valor]].copy()
    dados[coluna_valor] = pd.to_numeric(dados[coluna_valor], errors="coerce")
    dados = dados.dropna().sort_values(coluna_valor, ascending=False).head(top_n)
    dados = dados.sort_values(coluna_valor, ascending=True)

    fig, ax = plt.subplots(figsize=(9, max(4.5, 0.48 * len(dados) + 1.5)), dpi=180)
    barras = ax.barh(dados[coluna_categoria].astype(str), dados[coluna_valor])
    ax.set_title(titulo, fontsize=13, fontweight="bold", pad=12)
    ax.grid(axis="x", alpha=0.25)
    ax.set_axisbelow(True)

    for barra, valor in zip(barras, dados[coluna_valor]):
        texto = formatar_brl(valor) if rotulo_monetario else formatar_numero(valor)
        ax.text(
            barra.get_width(),
            barra.get_y() + barra.get_height() / 2,
            f"  {texto}",
            va="center",
            fontsize=8.5,
        )

    ax.tick_params(axis="y", labelsize=9)
    fig.tight_layout()
    fig.savefig(arquivo_saida, bbox_inches="tight")
    plt.close(fig)
    return arquivo_saida


def grafico_barras_empilhadas(
    df: pd.DataFrame,
    indice: str,
    coluna: str,
    valor: str,
    titulo: str,
    arquivo_saida: Path,
    top_n: int = 8,
) -> Path:
    """Gera barras empilhadas para participação de seguradoras por cultura."""
    dados = df[[indice, coluna, valor]].copy()
    dados[valor] = pd.to_numeric(dados[valor], errors="coerce").fillna(0)

    top_indice = (
        dados.groupby(indice)[valor]
        .sum()
        .sort_values(ascending=False)
        .head(top_n)
        .index
    )

    dados = dados[dados[indice].isin(top_indice)].copy()

    pivot = dados.pivot_table(
        index=indice,
        columns=coluna,
        values=valor,
        aggfunc="sum",
        fill_value=0,
    )

    # Mantém somente os 6 maiores concorrentes e agrupa o restante.
    top_colunas = pivot.sum().sort_values(ascending=False).head(6).index.tolist()
    outras = [c for c in pivot.columns if c not in top_colunas]

    pivot_plot = pivot[top_colunas].copy()
    if outras:
        pivot_plot["OUTRAS SEGURADORAS"] = pivot[outras].sum(axis=1)

    pivot_plot = pivot_plot.loc[
        pivot_plot.sum(axis=1).sort_values(ascending=False).index
    ]

    ax = pivot_plot.plot(
        kind="barh",
        stacked=True,
        figsize=(10, max(4.8, 0.6 * len(pivot_plot) + 1.6)),
        width=0.7,
    )

    ax.set_title(titulo, fontsize=13, fontweight="bold", pad=12)
    ax.set_xlabel("Prêmio líquido (R$)")
    ax.set_ylabel("")
    ax.grid(axis="x", alpha=0.25)
    ax.set_axisbelow(True)
    ax.legend(
        title="Seguradora",
        bbox_to_anchor=(1.02, 1),
        loc="upper left",
        fontsize=8,
    )

    fig = ax.get_figure()
    fig.tight_layout()
    fig.savefig(arquivo_saida, bbox_inches="tight")
    plt.close(fig)
    return arquivo_saida


def grafico_distribuicao_concentracao(
    df: pd.DataFrame,
    titulo: str,
    arquivo_saida: Path,
) -> Path:
    """Gera gráfico de distribuição de classes de concentração."""
    garantir_coluna(df, "classificacao_concentracao")

    ordem = [
        "BAIXA_CONCENTRACAO",
        "CONCENTRACAO_MODERADA",
        "ALTA_CONCENTRACAO",
    ]

    dados = (
        df["classificacao_concentracao"]
        .fillna("SEM INFORMACAO")
        .value_counts()
        .reindex(ordem, fill_value=0)
        .reset_index()
    )
    dados.columns = ["classe", "quantidade"]

    fig, ax = plt.subplots(figsize=(8, 4.8), dpi=180)
    barras = ax.bar(dados["classe"], dados["quantidade"])
    ax.set_title(titulo, fontsize=13, fontweight="bold", pad=12)
    ax.set_ylabel("Quantidade de hexágonos-cultura")
    ax.grid(axis="y", alpha=0.25)
    ax.set_axisbelow(True)
    ax.tick_params(axis="x", rotation=15)

    for barra, valor in zip(barras, dados["quantidade"]):
        ax.text(
            barra.get_x() + barra.get_width() / 2,
            barra.get_height(),
            formatar_numero(valor),
            ha="center",
            va="bottom",
            fontsize=9,
        )

    fig.tight_layout()
    fig.savefig(arquivo_saida, bbox_inches="tight")
    plt.close(fig)
    return arquivo_saida


# =============================================================================
# FUNÇÕES WORD
# =============================================================================


def set_cell_shading(cell, fill: str) -> None:
    """Aplica preenchimento a uma célula de tabela Word."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    """Define texto e estilo básico de uma célula."""
    cell.text = ""
    paragrafo = cell.paragraphs[0]
    run = paragrafo.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(8.5)
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def adicionar_titulo(doc: Document, texto: str, nivel: int = 1) -> None:
    """Adiciona título usando estilos do Word."""
    p = doc.add_paragraph()
    p.style = f"Heading {nivel}"
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(texto)
    run.font.color.rgb = RGBColor(31, 78, 121)


def adicionar_texto(doc: Document, texto: str, negrito_inicial: str | None = None) -> None:
    """Adiciona parágrafo justificado."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08

    if negrito_inicial and texto.startswith(negrito_inicial):
        p.add_run(negrito_inicial).bold = True
        p.add_run(texto[len(negrito_inicial):])
    else:
        p.add_run(texto)


def adicionar_imagem(doc: Document, caminho: Path, legenda: str) -> None:
    """Insere imagem centralizada com legenda."""
    if not caminho.exists():
        adicionar_texto(doc, f"[Imagem não gerada: {legenda}]")
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(caminho), width=LARGURA_IMAGEM_WORD)

    p_legenda = doc.add_paragraph()
    p_legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_legenda.paragraph_format.space_after = Pt(8)
    run = p_legenda.add_run(legenda)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(90, 90, 90)


def adicionar_tabela(
    doc: Document,
    df: pd.DataFrame,
    titulo: str,
    max_linhas: int = TOP_N,
) -> None:
    """Insere DataFrame simplificado como tabela Word."""
    adicionar_titulo(doc, titulo, nivel=2)

    if df.empty:
        adicionar_texto(doc, "Sem dados disponíveis para esta tabela.")
        return

    tabela_df = df.head(max_linhas).copy()
    tabela = doc.add_table(rows=1, cols=len(tabela_df.columns))
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.style = "Table Grid"

    for i, coluna in enumerate(tabela_df.columns):
        cell = tabela.rows[0].cells[i]
        set_cell_shading(cell, "1F4E79")
        set_cell_text(cell, coluna, bold=True, color=RGBColor(255, 255, 255))

    for _, linha in tabela_df.iterrows():
        celulas = tabela.add_row().cells
        for i, coluna in enumerate(tabela_df.columns):
            set_cell_text(celulas[i], linha[coluna])

    doc.add_paragraph()


def configurar_documento(doc: Document) -> None:
    """Aplica padrão visual ao documento."""
    secao = doc.sections[0]
    secao.top_margin = Inches(0.65)
    secao.bottom_margin = Inches(0.65)
    secao.left_margin = Inches(0.7)
    secao.right_margin = Inches(0.7)

    estilos = doc.styles
    estilos["Normal"].font.name = "Aptos"
    estilos["Normal"].font.size = Pt(10)
    estilos["Heading 1"].font.name = "Aptos Display"
    estilos["Heading 1"].font.size = Pt(18)
    estilos["Heading 2"].font.name = "Aptos Display"
    estilos["Heading 2"].font.size = Pt(13)


def adicionar_cabecalho_rodape(doc: Document) -> None:
    """Inclui cabeçalho e rodapé simples."""
    for secao in doc.sections:
        cabecalho = secao.header.paragraphs[0]
        cabecalho.text = "INTELIGÊNCIA TERRITORIAL | SEGURO RURAL"
        cabecalho.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in cabecalho.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(100, 100, 100)

        rodape = secao.footer.paragraphs[0]
        rodape.text = "Dados públicos do PSR | Período analisado: 2020–2024"
        rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in rodape.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(100, 100, 100)


def adicionar_card_indicadores(doc: Document, indicadores: list[tuple[str, str]]) -> None:
    """Cria uma tabela de cards para indicadores-chave."""
    colunas = min(3, len(indicadores))
    linhas = math.ceil(len(indicadores) / colunas)
    tabela = doc.add_table(rows=linhas, cols=colunas)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER

    indice = 0
    for linha in tabela.rows:
        for cell in linha.cells:
            if indice < len(indicadores):
                titulo, valor = indicadores[indice]
                set_cell_shading(cell, "EAF1F8")
                cell.text = ""
                p1 = cell.paragraphs[0]
                p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r1 = p1.add_run(valor)
                r1.bold = True
                r1.font.size = Pt(14)
                r1.font.color.rgb = RGBColor(31, 78, 121)
                p2 = cell.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r2 = p2.add_run(titulo)
                r2.font.size = Pt(8)
                r2.font.color.rgb = RGBColor(80, 80, 80)
            indice += 1

    doc.add_paragraph()


# =============================================================================
# PRINCIPAL
# =============================================================================


def main() -> None:
    """Executa geração de imagens e relatório Word."""
    print("=" * 90)
    print("GERAÇÃO AUTOMÁTICA DE RELATÓRIO PSR")
    print("=" * 90)

    verificar_arquivo(GPKG_TOTAL, obrigatorio=True)
    verificar_arquivo(GPKG_CULTURA, obrigatorio=True)
    tem_sem_dominantes = verificar_arquivo(GPKG_SEM_DOMINANTES, obrigatorio=False)

    if LIMPAR_IMAGENS_ANTERIORES:
        for arquivo in PASTA_IMAGENS.glob("*.png"):
            arquivo.unlink()
        for arquivo in PASTA_TABELAS.glob("*.csv"):
            arquivo.unlink()
        if ARQUIVO_DOCX.exists():
            ARQUIVO_DOCX.unlink()

    contorno = carregar_contorno()

    camada_total = f"{PREFIXO_GRADE}_acum_2020_2024"
    camada_cult = f"{PREFIXO_GRADE}_acum_cult"
    camada_cultseg = f"{PREFIXO_GRADE}_acum_cultseg"

    print(f"Lendo camada total: {camada_total}")
    gdf_total = ler_camada(GPKG_TOTAL, camada_total)

    print(f"Lendo camada por cultura: {camada_cult}")
    gdf_cult = ler_camada(GPKG_CULTURA, camada_cult)

    print(f"Lendo camada por cultura e seguradora: {camada_cultseg}")
    gdf_cultseg = ler_camada(GPKG_CULTURA, camada_cultseg)

    # Padroniza culturas e reagrupa MILHO CONSORCIADO como MILHO 2ª SAFRA.
    # Isso também recalcula líder, share e HHI para a cultura combinada.
    gdf_cult, gdf_cultseg = reagregar_camada_cultura_para_relatorio(
        gdf_cult,
        gdf_cultseg,
    )

    # -------------------------------------------------------------------------
    # Tabelas analíticas de apoio
    # -------------------------------------------------------------------------
    for coluna in [
        "premio_liq_rs",
        "limite_garantia_rs",
        "indenizacao_rs",
        "area_total_ha",
    ]:
        if coluna in gdf_total.columns:
            gdf_total[coluna] = pd.to_numeric(gdf_total[coluna], errors="coerce").fillna(0)

    for coluna in [
        "premio_liq_rs",
        "limite_garantia_rs",
        "indenizacao_rs",
        "area_total_ha",
        "hhi_premio",
        "share_seguradora_lider_pct",
    ]:
        if coluna in gdf_cult.columns:
            gdf_cult[coluna] = pd.to_numeric(gdf_cult[coluna], errors="coerce")

    for coluna in [
        "premio_liq_rs",
        "limite_garantia_rs",
        "indenizacao_rs",
        "area_total_ha",
        "market_share_seg_pct",
    ]:
        if coluna in gdf_cultseg.columns:
            gdf_cultseg[coluna] = pd.to_numeric(gdf_cultseg[coluna], errors="coerce")

    resumo_culturas = (
        gdf_cult[gdf_cult["cultura"].isin(CULTURAS_DESTAQUE)]
        .groupby("cultura", as_index=False)
        .agg(
            premio_liq_rs=("premio_liq_rs", "sum"),
            limite_garantia_rs=("limite_garantia_rs", "sum"),
            area_total_ha=("area_total_ha", "sum"),
            indenizacao_rs=("indenizacao_rs", "sum"),
            hexagonos=("grid_id", "nunique"),
            seguradoras=("n_seguradoras", "sum"),
        )
        .sort_values("premio_liq_rs", ascending=False)
    )

    ranking_seguradoras = (
        gdf_cultseg[gdf_cultseg["cultura"].isin(CULTURAS_DESTAQUE)]
        .groupby("seguradora", as_index=False)
        .agg(
            premio_liq_rs=("premio_liq_rs", "sum"),
            limite_garantia_rs=("limite_garantia_rs", "sum"),
            indenizacao_rs=("indenizacao_rs", "sum"),
            culturas=("cultura", "nunique"),
            hexagonos=("grid_id", "nunique"),
        )
        .sort_values("premio_liq_rs", ascending=False)
    )

    top_hexagonos = (
        gdf_total[["grid_id", "premio_liq_rs", "limite_garantia_rs", "indenizacao_rs", "n_registros"]]
        .sort_values("premio_liq_rs", ascending=False)
        .head(TOP_N)
        .copy()
    )

    # -------------------------------------------------------------------------
    # Mapas e gráficos
    # -------------------------------------------------------------------------
    imagens: dict[str, Path] = {}

    imagens["mapa_premio_total"] = mapa_graduado(
        gdf_total,
        "premio_liq_rs",
        "Mercado segurado: prêmio líquido acumulado (2020–2024)",
        PASTA_IMAGENS / "01_mapa_premio_total.png",
        contorno,
        "Prêmio líquido (R$)",
    )

    imagens["mapa_limite_total"] = mapa_graduado(
        gdf_total,
        "limite_garantia_rs",
        "Exposição financeira: limite de garantia acumulado (2020–2024)",
        PASTA_IMAGENS / "02_mapa_limite_total.png",
        contorno,
        "Limite de garantia (R$)",
    )

    imagens["grafico_culturas"] = grafico_barras_horizontais(
        resumo_culturas,
        "cultura",
        "premio_liq_rs",
        "Prêmio líquido acumulado por cultura",
        PASTA_IMAGENS / "03_grafico_premio_culturas.png",
        top_n=len(resumo_culturas),
        rotulo_monetario=True,
    )

    imagens["grafico_seguradoras"] = grafico_barras_horizontais(
        ranking_seguradoras,
        "seguradora",
        "premio_liq_rs",
        "Principais seguradoras por prêmio nas culturas selecionadas",
        PASTA_IMAGENS / "04_grafico_ranking_seguradoras.png",
        top_n=TOP_N,
        rotulo_monetario=True,
    )

    imagens["grafico_mix"] = grafico_barras_empilhadas(
        gdf_cultseg[gdf_cultseg["cultura"].isin(CULTURAS_DESTAQUE)],
        "cultura",
        "seguradora",
        "premio_liq_rs",
        "Mix de seguradoras por cultura",
        PASTA_IMAGENS / "05_grafico_mix_seguradoras_culturas.png",
        top_n=len(CULTURAS_DESTAQUE),
    )

    gdf_concorrencia = gdf_cult[
        gdf_cult["cultura"].eq(CULTURA_MAPA_CONCORRENCIA)
    ].copy()

    if not gdf_concorrencia.empty:
        nome_cult = limpar_nome_arquivo(CULTURA_MAPA_CONCORRENCIA)

        imagens["mapa_cultura"] = mapa_graduado(
            gdf_concorrencia,
            "premio_liq_rs",
            f"{CULTURA_MAPA_CONCORRENCIA}: prêmio líquido por hexágono (2020–2024)",
            PASTA_IMAGENS / f"06_mapa_{nome_cult}_premio.png",
            contorno,
            "Prêmio líquido (R$)",
        )

        imagens["mapa_lider"] = mapa_categorico(
            gdf_concorrencia,
            "seguradora_lider",
            f"{CULTURA_MAPA_CONCORRENCIA}: seguradora líder por hexágono",
            PASTA_IMAGENS / f"07_mapa_{nome_cult}_lider.png",
            contorno,
            max_categorias=11,
        )

        imagens["mapa_hhi"] = mapa_graduado(
            gdf_concorrencia,
            "hhi_premio",
            f"{CULTURA_MAPA_CONCORRENCIA}: concentração competitiva (HHI)",
            PASTA_IMAGENS / f"08_mapa_{nome_cult}_hhi.png",
            contorno,
            "HHI de prêmio",
        )

        imagens["grafico_concentracao"] = grafico_distribuicao_concentracao(
            gdf_concorrencia,
            f"{CULTURA_MAPA_CONCORRENCIA}: distribuição de concentração por hexágono",
            PASTA_IMAGENS / f"09_grafico_{nome_cult}_concentracao.png",
        )

    # Mapas das culturas selecionadas para usar no relatório/anexo.
    for cultura in CULTURAS_DESTAQUE:
        subset = gdf_cult[gdf_cult["cultura"].eq(cultura)].copy()
        if subset.empty:
            continue
        chave = f"mapa_{limpar_nome_arquivo(cultura)}"
        imagens[chave] = mapa_graduado(
            subset,
            "premio_liq_rs",
            f"{cultura}: prêmio líquido por hexágono (2020–2024)",
            PASTA_IMAGENS / f"mapa_{limpar_nome_arquivo(cultura)}.png",
            contorno,
            "Prêmio líquido (R$)",
        )

    # Cenário sem dominantes (caso o GPKG esteja disponível).
    gdf_sem_dom = None
    if tem_sem_dominantes:
        camada_sem_dom = f"{PREFIXO_GRADE}_acum_cult_sem_dom"
        try:
            gdf_sem_dom = ler_camada(GPKG_SEM_DOMINANTES, camada_sem_dom)
            subset_sem_dom = gdf_sem_dom[
                gdf_sem_dom["cultura"].eq(CULTURA_MAPA_CONCORRENCIA)
            ].copy()

            if not subset_sem_dom.empty:
                imagens["mapa_sem_dominantes"] = mapa_categorico(
                    subset_sem_dom,
                    "seguradora_lider",
                    (
                        f"{CULTURA_MAPA_CONCORRENCIA}: líder sem seguradoras dominantes "
                        "(mercado remanescente)"
                    ),
                    PASTA_IMAGENS / "10_mapa_sem_dominantes_lider.png",
                    contorno,
                    max_categorias=11,
                )
        except Exception as exc:
            print(f"Aviso: cenário sem dominantes não foi incluído: {exc}")

    # Exporta as tabelas para CSV, inclusive para eventual revisão antes do Word.
    resumo_culturas.to_csv(
        PASTA_TABELAS / "resumo_culturas.csv",
        index=False,
        encoding="utf-8-sig",
    )
    ranking_seguradoras.to_csv(
        PASTA_TABELAS / "ranking_seguradoras.csv",
        index=False,
        encoding="utf-8-sig",
    )
    top_hexagonos.to_csv(
        PASTA_TABELAS / "top_hexagonos_premio.csv",
        index=False,
        encoding="utf-8-sig",
    )

    # -------------------------------------------------------------------------
    # Métricas executivas
    # -------------------------------------------------------------------------
    premio_total = gdf_total["premio_liq_rs"].sum()
    limite_total = gdf_total["limite_garantia_rs"].sum()
    indenizacao_total = gdf_total["indenizacao_rs"].sum()
    n_hex = gdf_total["grid_id"].nunique()
    n_registros = int(pd.to_numeric(gdf_total.get("n_registros", 0), errors="coerce").fillna(0).sum())
    n_culturas = gdf_cult["cultura"].nunique()
    n_seguradoras = gdf_cultseg["seguradora"].nunique()

    # -------------------------------------------------------------------------
    # GERAÇÃO DO WORD
    # -------------------------------------------------------------------------
    print("Gerando relatório Word...")
    doc = Document()
    configurar_documento(doc)
    adicionar_cabecalho_rodape(doc)

    # Capa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("INTELIGÊNCIA TERRITORIAL\nDO MERCADO DE SEGURO RURAL")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(31, 78, 121)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Programa de Subvenção ao Prêmio do Seguro Rural (PSR)\n2020–2024")
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(80, 80, 80)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(26)
    r3 = p3.add_run(
        "Relatório gerado automaticamente a partir das camadas vetoriais do projeto"
    )
    r3.italic = True
    r3.font.size = Pt(10)

    doc.add_page_break()

    # Resumo executivo
    adicionar_titulo(doc, "1. Resumo executivo", nivel=1)
    adicionar_texto(
        doc,
        "Este relatório consolida uma visão territorial do mercado de seguro rural a partir das camadas vetoriais geradas com dados públicos do PSR. A análise combina grade hexagonal, cultura e seguradora, permitindo identificar concentração de prêmio, exposição financeira, presença competitiva, liderança territorial e valores de indenização registrados.",
    )

    adicionar_card_indicadores(
        doc,
        [
            ("Prêmio líquido na grade", formatar_brl(premio_total)),
            ("Limite de garantia na grade", formatar_brl(limite_total)),
            ("Indenização registrada", formatar_brl(indenizacao_total)),
            ("Hexágonos com registros", formatar_numero(n_hex)),
            ("Registros agregados", formatar_numero(n_registros)),
            ("Seguradoras identificadas", formatar_numero(n_seguradoras)),
        ],
    )

    adicionar_texto(
        doc,
        "Leitura executiva: a solução permite passar de um ranking nacional ou municipal para uma visão espacial que mostra onde cada cultura se concentra, quais seguradoras atuam em cada território e quais áreas apresentam maior exposição ou concentração competitiva.",
    )

    # Metodologia
    adicionar_titulo(doc, "2. Base de dados e abordagem", nivel=1)
    adicionar_texto(
        doc,
        "A estrutura analítica combina três escalas: município para leitura agregada do mercado, pontos georreferenciados para observar a distribuição espacial dos registros e hexágonos para medir acúmulo, concorrência e perfil de risco em uma unidade territorial comparável.",
    )

    tabela_metodo = pd.DataFrame(
        [
            ["Município", "Mercado agregado", "Prêmio, limite, cultura e seguradora líder"],
            ["Ponto georreferenciado", "Localização informada", "Distribuição e recorrência espacial dos registros"],
            ["Hexágono", "Unidade territorial de análise", "Concentração, concorrência, exposição e indenização"],
        ],
        columns=["Camada", "Unidade de análise", "Uso principal"],
    )
    adicionar_tabela(doc, tabela_metodo, "Estrutura territorial da análise", max_linhas=10)

    adicionar_texto(
        doc,
        "Limitação importante: os pontos representam coordenadas informadas na base PSR e não o limite real das propriedades. Além disso, os valores de indenização devem ser interpretados como indenizações registradas, e não como sinistralidade técnica definitiva.",
    )

    # Mercado geral
    adicionar_titulo(doc, "3. Mercado e exposição territorial", nivel=1)
    adicionar_texto(
        doc,
        "Os dois mapas a seguir mostram, respectivamente, onde se concentra o prêmio líquido e onde se concentra o limite de garantia. A comparação entre eles ajuda a distinguir territórios de maior atividade comercial daqueles com maior acúmulo financeiro potencial.",
    )
    adicionar_imagem(doc, imagens["mapa_premio_total"], "Figura 1. Prêmio líquido acumulado por hexágono.")
    adicionar_imagem(doc, imagens["mapa_limite_total"], "Figura 2. Limite de garantia acumulado por hexágono.")

    # Culturas
    adicionar_titulo(doc, "4. Perfil das culturas", nivel=1)
    adicionar_texto(
        doc,
        "A análise por cultura permite identificar quais atividades sustentam o mercado segurado e quais representam nichos territorialmente concentrados. Os valores abaixo correspondem às culturas selecionadas no script de processamento.",
    )
    adicionar_imagem(doc, imagens["grafico_culturas"], "Figura 3. Prêmio líquido acumulado por cultura.")

    tabela_culturas_word = resumo_culturas.copy()
    tabela_culturas_word["premio_liq_rs"] = tabela_culturas_word["premio_liq_rs"].map(formatar_brl)
    tabela_culturas_word["limite_garantia_rs"] = tabela_culturas_word["limite_garantia_rs"].map(formatar_brl)
    tabela_culturas_word["area_total_ha"] = tabela_culturas_word["area_total_ha"].map(lambda x: formatar_numero(x, 0))
    tabela_culturas_word["indenizacao_rs"] = tabela_culturas_word["indenizacao_rs"].map(formatar_brl)
    tabela_culturas_word = tabela_culturas_word.rename(
        columns={
            "cultura": "Cultura",
            "premio_liq_rs": "Prêmio líquido",
            "limite_garantia_rs": "Limite de garantia",
            "area_total_ha": "Área (ha)",
            "indenizacao_rs": "Indenização",
            "hexagonos": "Hexágonos",
            "seguradoras": "Soma de seguradoras",
        }
    )
    adicionar_tabela(doc, tabela_culturas_word, "Resumo das culturas selecionadas", max_linhas=len(tabela_culturas_word))

    # Concorrência
    adicionar_titulo(doc, "5. Concorrência e seguradoras", nivel=1)
    adicionar_texto(
        doc,
        "A camada hexágono + cultura + seguradora permite distinguir presença comercial de liderança efetiva. A análise mostra quais seguradoras concentram prêmio nas culturas selecionadas e como o mix competitivo varia por atividade.",
    )
    adicionar_imagem(doc, imagens["grafico_seguradoras"], "Figura 4. Ranking de seguradoras por prêmio nas culturas selecionadas.")
    adicionar_imagem(doc, imagens["grafico_mix"], "Figura 5. Mix de seguradoras por cultura.")

    tabela_seg_word = ranking_seguradoras.head(TOP_N).copy()
    tabela_seg_word["premio_liq_rs"] = tabela_seg_word["premio_liq_rs"].map(formatar_brl)
    tabela_seg_word["limite_garantia_rs"] = tabela_seg_word["limite_garantia_rs"].map(formatar_brl)
    tabela_seg_word["indenizacao_rs"] = tabela_seg_word["indenizacao_rs"].map(formatar_brl)
    tabela_seg_word = tabela_seg_word.rename(
        columns={
            "seguradora": "Seguradora",
            "premio_liq_rs": "Prêmio líquido",
            "limite_garantia_rs": "Limite de garantia",
            "indenizacao_rs": "Indenização",
            "culturas": "Culturas",
            "hexagonos": "Hexágonos",
        }
    )
    adicionar_tabela(doc, tabela_seg_word, "Principais seguradoras", max_linhas=TOP_N)

    # Concorrência da cultura escolhida
    if not gdf_concorrencia.empty:
        adicionar_titulo(doc, f"6. Concorrência territorial: {CULTURA_MAPA_CONCORRENCIA}", nivel=1)
        adicionar_texto(
            doc,
            f"Os mapas abaixo mostram onde o mercado de {CULTURA_MAPA_CONCORRENCIA} se concentra, quem lidera em cada hexágono e como varia a concentração competitiva. O HHI mede se o prêmio está distribuído entre várias seguradoras ou concentrado em poucas.",
        )
        adicionar_imagem(doc, imagens["mapa_cultura"], f"Figura 6. {CULTURA_MAPA_CONCORRENCIA}: prêmio líquido por hexágono.")
        adicionar_imagem(doc, imagens["mapa_lider"], f"Figura 7. {CULTURA_MAPA_CONCORRENCIA}: seguradora líder por hexágono.")
        adicionar_imagem(doc, imagens["mapa_hhi"], f"Figura 8. {CULTURA_MAPA_CONCORRENCIA}: HHI de prêmio por hexágono.")
        adicionar_imagem(doc, imagens["grafico_concentracao"], f"Figura 9. {CULTURA_MAPA_CONCORRENCIA}: distribuição das classes de concentração.")

    # Sem dominantes
    adicionar_titulo(doc, "7. Mercado remanescente sem seguradoras dominantes", nivel=1)
    if "mapa_sem_dominantes" in imagens:
        adicionar_texto(
            doc,
            "A leitura sem seguradoras dominantes é útil para revelar a concorrência entre as demais empresas. Prêmio, participação de mercado, liderança e HHI são recalculados considerando apenas o mercado remanescente, e não apenas filtrados visualmente no QGIS.",
        )
        adicionar_imagem(doc, imagens["mapa_sem_dominantes"], "Figura 10. Liderança no mercado remanescente, após exclusão das seguradoras dominantes definidas no script.")
    else:
        adicionar_texto(
            doc,
            "O GeoPackage de mercado sem seguradoras dominantes não foi localizado ou não continha a camada esperada. Quando o script 07d estiver executado, esta seção será preenchida automaticamente.",
        )

    # Roteiro QGIS
    adicionar_titulo(doc, "8. Roteiro de exploração ao vivo no QGIS", nivel=1)
    adicionar_texto(
        doc,
        "O relatório resume os resultados. Para responder perguntas específicas durante a reunião, o QGIS permite filtrar cultura, seguradora, ano e escala espacial. A tabela abaixo sugere a camada e a simbologia mais adequada para cada pergunta executiva.",
    )

    roteiro_qgis = pd.DataFrame(
        [
            ["Onde está o maior mercado?", f"{PREFIXO_GRADE}_acum_2020_2024", "premio_liq_rs", "Graduado (5–6 classes)"],
            ["Onde há maior exposição?", f"{PREFIXO_GRADE}_acum_2020_2024", "limite_garantia_rs", "Graduado (5–6 classes)"],
            ["Onde está uma cultura?", f"{PREFIXO_GRADE}_acum_cult", "premio_liq_rs", "Filtro por cultura + graduado"],
            ["Quem lidera uma cultura?", f"{PREFIXO_GRADE}_acum_cult", "seguradora_lider", "Categorizado"],
            ["Quais seguradoras atuam?", f"{PREFIXO_GRADE}_acum_cultseg", "seguradora", "Filtro por cultura + categorizado"],
            ["Onde há concorrência?", f"{PREFIXO_GRADE}_acum_cult", "classificacao_concentracao", "Categorizado"],
            ["Quem lidera sem dominantes?", f"{PREFIXO_GRADE}_acum_cult_sem_dom", "seguradora_lider", "Categorizado"],
            ["Onde há indenização registrada?", f"{PREFIXO_GRADE}_acum_2020_2024", "indenizacao_rs", "Graduado"],
        ],
        columns=["Pergunta", "Camada", "Campo", "Simbologia"],
    )
    adicionar_tabela(doc, roteiro_qgis, "Camadas e simbologias recomendadas", max_linhas=20)

    # Oportunidade
    adicionar_titulo(doc, "9. Aplicações estratégicas e próximos passos", nivel=1)
    adicionar_texto(
        doc,
        "A solução pode apoiar expansão comercial, priorização territorial, acompanhamento de concorrência, identificação de concentração de exposição e seleção de áreas para aprofundamento técnico. O próximo passo natural é incluir a carteira interna, permitindo calcular participação própria, sobreposição competitiva e um score de oportunidade por hexágono e cultura.",
    )

    sugestoes = pd.DataFrame(
        [
            ["Expansão comercial", "Prêmio alto + HHI baixo/moderado + baixa participação própria"],
            ["Concorrência", "Líder, share da líder e HHI por hexágono + cultura"],
            ["Acúmulo", "Limite de garantia, área e número de registros por hexágono"],
            ["Risco registrado", "Indenização, frequência de indenização e evento líder"],
            ["Mercado acessível", "Comparação entre mercado total e mercado sem dominantes"],
        ],
        columns=["Aplicação", "Indicadores prioritários"],
    )
    adicionar_tabela(doc, sugestoes, "Aplicações práticas", max_linhas=10)

    # Apêndice com mapas por cultura
    doc.add_page_break()
    adicionar_titulo(doc, "Apêndice A. Mapas por cultura", nivel=1)
    adicionar_texto(
        doc,
        "Os mapas abaixo são gerados diretamente a partir das camadas de hexágonos por cultura. Eles podem ser substituídos ou complementados por exports do QGIS caso seja necessário aplicar identidade visual corporativa, rótulos, escala gráfica e logotipo.",
    )

    for cultura in CULTURAS_DESTAQUE:
        chave = f"mapa_{limpar_nome_arquivo(cultura)}"
        if chave in imagens:
            adicionar_titulo(doc, cultura, nivel=2)
            adicionar_imagem(doc, imagens[chave], f"Mapa de prêmio líquido acumulado — {cultura}.")

    # Salva
    doc.save(ARQUIVO_DOCX)

    print("=" * 90)
    print("RELATÓRIO GERADO COM SUCESSO")
    print("=" * 90)
    print(f"Word: {ARQUIVO_DOCX}")
    print(f"Imagens: {PASTA_IMAGENS}")
    print(f"Tabelas: {PASTA_TABELAS}")


if __name__ == "__main__":
    main()
